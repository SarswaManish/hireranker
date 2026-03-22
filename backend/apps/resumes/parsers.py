import logging
import os
import re
import tempfile
from typing import Optional
from urllib.parse import urlparse

from django.utils import timezone

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using pdfplumber.

    Args:
        file_path: Absolute or Django storage path to the PDF

    Returns:
        Extracted text string
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is not installed. Run: pip install pdfplumber")

    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning("Failed to extract page %d from PDF %s: %s", page_num, file_path, e)

        full_text = '\n\n'.join(text_parts)
        # Clean up excessive whitespace
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)
        full_text = re.sub(r'[ \t]+', ' ', full_text)
        return full_text.strip()

    except Exception as e:
        logger.error("Failed to extract text from PDF %s: %s", file_path, e)
        raise ValueError(f"PDF extraction failed: {str(e)}") from e


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text string
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    try:
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(' | '.join(row_texts))

        full_text = '\n'.join(text_parts)
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)
        return full_text.strip()

    except Exception as e:
        logger.error("Failed to extract text from DOCX %s: %s", file_path, e)
        raise ValueError(f"DOCX extraction failed: {str(e)}") from e


def extract_text_from_url(url: str, timeout: int = 30) -> str:
    """
    Fetch a resume from URL and extract text.
    Supports PDF and DOCX URLs, and Google Drive/Docs links.

    Args:
        url: URL to the resume
        timeout: Request timeout in seconds

    Returns:
        Extracted text string
    """
    import requests

    # Handle Google Drive links
    if 'drive.google.com' in url:
        url = _convert_gdrive_url(url)

    headers = {
        'User-Agent': 'Mozilla/5.0 (compatible; HireRanker/1.0)',
    }

    try:
        response = requests.get(url, headers=headers, timeout=timeout, stream=True)
        response.raise_for_status()

        content_type = response.headers.get('content-type', '').lower()
        content_length = int(response.headers.get('content-length', 0))

        max_size = 10 * 1024 * 1024  # 10MB
        if content_length > max_size:
            raise ValueError(f"File too large: {content_length} bytes")

        content = b''
        for chunk in response.iter_content(chunk_size=8192):
            content += chunk
            if len(content) > max_size:
                raise ValueError("File too large (exceeds 10MB)")

        # Determine file type
        if 'pdf' in content_type or url.lower().endswith('.pdf') or content[:4] == b'%PDF':
            return _extract_from_bytes(content, 'pdf')
        elif (
            'officedocument' in content_type
            or 'msword' in content_type
            or url.lower().endswith('.docx')
        ):
            return _extract_from_bytes(content, 'docx')
        else:
            # Try PDF first, then DOCX
            if content[:4] == b'%PDF':
                return _extract_from_bytes(content, 'pdf')
            elif content[:2] == b'PK':  # ZIP/DOCX signature
                return _extract_from_bytes(content, 'docx')
            else:
                # Plain text fallback
                try:
                    return content.decode('utf-8', errors='replace').strip()
                except Exception:
                    raise ValueError("Unable to determine or parse file type from URL")

    except requests.RequestException as e:
        logger.error("Failed to fetch URL %s: %s", url, e)
        raise ValueError(f"Failed to fetch resume URL: {str(e)}") from e


def _convert_gdrive_url(url: str) -> str:
    """Convert Google Drive share URL to direct download URL."""
    # Pattern: https://drive.google.com/file/d/{FILE_ID}/view
    match = re.search(r'/file/d/([a-zA-Z0-9_-]+)', url)
    if match:
        file_id = match.group(1)
        return f"https://drive.google.com/uc?export=download&id={file_id}"
    # Pattern: https://drive.google.com/open?id={FILE_ID}
    match = re.search(r'[?&]id=([a-zA-Z0-9_-]+)', url)
    if match:
        file_id = match.group(1)
        return f"https://drive.google.com/uc?export=download&id={file_id}"
    return url


def _extract_from_bytes(content: bytes, file_type: str) -> str:
    """Extract text from bytes by writing to temp file."""
    suffix = '.pdf' if file_type == 'pdf' else '.docx'
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_file:
        tmp_file.write(content)
        tmp_path = tmp_file.name

    try:
        if file_type == 'pdf':
            return extract_text_from_pdf(tmp_path)
        else:
            return extract_text_from_docx(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def parse_resume(resume) -> None:
    """
    Parse a CandidateResume and update it with extracted text.

    Args:
        resume: CandidateResume instance (from apps.candidates.models)
    """
    from apps.candidates.models import CandidateResume

    resume.parse_status = CandidateResume.ParseStatus.PARSING
    resume.save(update_fields=['parse_status'])

    try:
        raw_text = ''

        if resume.file_type == CandidateResume.FileType.PDF and resume.file_path:
            try:
                file_path = resume.file_path.path
            except Exception:
                file_path = str(resume.file_path)
            raw_text = extract_text_from_pdf(file_path)

        elif resume.file_type == CandidateResume.FileType.DOCX and resume.file_path:
            try:
                file_path = resume.file_path.path
            except Exception:
                file_path = str(resume.file_path)
            raw_text = extract_text_from_docx(file_path)

        elif resume.file_type == CandidateResume.FileType.URL:
            # Use the candidate's resume_url
            resume_url = resume.candidate.resume_url
            if not resume_url:
                raise ValueError("No URL provided for URL-type resume")
            raw_text = extract_text_from_url(resume_url)

        if not raw_text:
            raise ValueError("No text could be extracted from the resume")

        resume.raw_text = raw_text
        resume.parse_status = CandidateResume.ParseStatus.COMPLETED
        resume.parsed_at = timezone.now()
        resume.parse_error = ''
        resume.save(update_fields=['raw_text', 'parse_status', 'parsed_at', 'parse_error'])

        logger.info(
            "Resume parsed for candidate %s: %d chars extracted",
            resume.candidate_id, len(raw_text)
        )

    except Exception as e:
        logger.error("Resume parsing failed for candidate %s: %s", resume.candidate_id, e)
        resume.parse_status = CandidateResume.ParseStatus.FAILED
        resume.parse_error = str(e)
        resume.save(update_fields=['parse_status', 'parse_error'])
        raise
